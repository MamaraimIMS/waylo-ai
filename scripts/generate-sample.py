#!/usr/bin/env python3
"""
WayloAI · Test Document Generator

Генерирует реалистичные тестовые заявки DMC в формате DOCX для тестирования
AI-парсинга. Используется в demo, E2E тестах и QA.

Usage:
    python3 generate-sample.py
    python3 generate-sample.py --output custom.docx --days 7

Requires:
    pip install python-docx --break-system-packages
"""

import argparse
from datetime import date, timedelta
from pathlib import Path
import random
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Missing dependency. Install with:")
    print("  pip install python-docx --break-system-packages")
    sys.exit(1)


SAMPLE_TOURS = [
    {
        "client": "Hugo Villar",
        "country": "USA",
        "pax": 2,
        "days": 10,
        "cities": ["Tashkent", "Khiva", "Bukhara", "Samarkand"],
    },
    {
        "client": "Sarah Chen",
        "country": "UK",
        "pax": 1,
        "days": 6,
        "cities": ["Tashkent", "Samarkand", "Bukhara"],
    },
    {
        "client": "The Rodriguez Family",
        "country": "Spain",
        "pax": 4,
        "days": 8,
        "cities": ["Tashkent", "Khiva", "Bukhara", "Samarkand"],
    },
    {
        "client": "Martin Petersen",
        "country": "Germany",
        "pax": 2,
        "days": 7,
        "cities": ["Tashkent", "Samarkand", "Bukhara"],
    },
]

TRANSFER_TEMPLATES = {
    "arrival": ("{city} — Airport", "{city} — Hotel", 20, 35),
    "departure": ("{city} — Hotel", "{city} — Airport", 20, 40),
    "station_pickup": ("{city} — Railway", "{city} — Hotel", 20, 30),
    "station_dropoff": ("{city} — Hotel", "{city} — Railway", 20, 30),
    "restaurant": ("{city} — Hotel", "{city} — Restaurant", 15, 40),
    "hotel_return": ("{city} — Restaurant", "{city} — Hotel", 15, 40),
    "excursion": ("{city} — Hotel", "{city} — Sightseeing area", 30, 60),
    "excursion_return": ("{city} — Sightseeing area", "{city} — Hotel", 30, 60),
}


def generate_tour_program(tour):
    """Generate a realistic tour program as text suitable for DOCX"""
    client = tour["client"]
    country = tour["country"]
    pax = tour["pax"]
    days = tour["days"]
    cities = tour["cities"]

    start_date = date(2026, 5, 10) + timedelta(days=random.randint(0, 60))

    lines = []
    lines.append(f"TOUR PROGRAM: {client} ({country})")
    lines.append(f"Pax: {pax}")
    lines.append(f"Duration: {days} days / {days-1} nights")
    lines.append(f"Start date: {start_date.strftime('%d.%m.%Y')}")
    lines.append(f"Vehicle: Sedan (economy) for 1-2 pax" if pax <= 2 else f"Minivan for {pax} pax")
    lines.append("")
    lines.append("ITINERARY:")
    lines.append("=" * 50)

    # Distribute days across cities (first city gets arrival, last gets departure)
    days_per_city = days // len(cities)
    extra = days % len(cities)

    current_day = 1
    current_city = 0
    for ci, city in enumerate(cities):
        city_days = days_per_city + (1 if ci < extra else 0)

        for day_in_city in range(city_days):
            current_date = start_date + timedelta(days=current_day - 1)
            lines.append("")
            lines.append(f"Day {current_day}. {current_date.strftime('%d.%m.%Y')} · {city}")
            lines.append("-" * 40)

            transfers = []

            # First day of first city: arrival
            if current_day == 1 and ci == 0:
                transfers.append((("10:00", "10:20"), "Arrival at airport", "Airport", "Hotel"))
                transfers.append((("13:00", "15:00"), "City tour", "Hotel", "Old town"))
                transfers.append((("15:00", "17:00"), "Return to hotel", "Old town", "Hotel"))

            # Last day of last city: departure
            elif current_day == days and ci == len(cities) - 1:
                transfers.append((("09:00", "09:30"), "Check-out and transfer to airport", "Hotel", "Airport"))

            # Inter-city transfer day
            elif day_in_city == 0 and ci > 0:
                prev_city = cities[ci - 1]
                if random.random() > 0.5:
                    # Train transfer
                    transfers.append((("08:00", "08:30"), f"Transfer to {prev_city} railway station", f"{prev_city} — Hotel", f"{prev_city} — Railway"))
                    transfers.append((("14:00", "14:30"), f"Arrival at {city}, transfer to hotel", f"{city} — Railway", f"{city} — Hotel"))
                else:
                    # Flight transfer
                    transfers.append((("08:00", "08:30"), f"Transfer to {prev_city} airport", f"{prev_city} — Hotel", f"{prev_city} — Airport"))
                    transfers.append((("12:00", "12:35"), f"Arrival at {city}", f"{city} — Airport", f"{city} — Hotel"))

            # Regular sightseeing day
            else:
                transfers.append((("09:00", "13:00"), f"Morning excursion", f"{city} — Hotel", f"{city} — Sightseeing"))
                transfers.append((("13:00", "14:00"), "Lunch at restaurant", f"{city} — Sightseeing", f"{city} — Restaurant"))
                transfers.append((("14:00", "18:00"), "Afternoon program", f"{city} — Restaurant", f"{city} — Hotel"))

            for (times, desc, from_loc, to_loc) in transfers:
                lines.append(f"  {times[0]} - {times[1]}: {desc}")
                lines.append(f"    From: {from_loc}")
                lines.append(f"    To: {to_loc}")
                lines.append("")

            current_day += 1

            if current_day > days:
                break
        if current_day > days:
            break

    lines.append("")
    lines.append("=" * 50)
    lines.append("END OF PROGRAM")
    lines.append("")
    lines.append("Notes:")
    lines.append("- All transfers by licensed drivers")
    lines.append("- English speaking guides required")
    lines.append("- Please confirm availability 48h before arrival")

    return "\n".join(lines)


def generate_docx(tour, output_path):
    """Generate a .docx file from tour program"""
    doc = Document()

    # Title
    title = doc.add_heading(f'{tour["client"]}, {tour["country"]}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f'Private Tour · {tour["days"]} days / {tour["days"]-1} nights · {tour["pax"]} pax')
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(0x72, 0x72, 0x6C)

    doc.add_paragraph()

    # Program text
    program_text = generate_tour_program(tour)
    for line in program_text.split('\n'):
        p = doc.add_paragraph()
        if line.startswith("Day ") or line.startswith("TOUR PROGRAM") or line.startswith("ITINERARY"):
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
        elif line.startswith("="):
            run = p.add_run(line)
            run.font.color.rgb = RGBColor(0x0A, 0x7D, 0x5F)
        else:
            p.add_run(line)
            p.runs[0].font.size = Pt(10)

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description='Generate sample DMC request DOCX for testing WayloAI')
    parser.add_argument('--output', default='test-request.docx', help='Output file path')
    parser.add_argument('--tour', default=None, help='Specific tour name (Hugo Villar, Sarah Chen, etc.)')
    parser.add_argument('--random', action='store_true', help='Generate random tour from templates')
    args = parser.parse_args()

    # Pick tour
    if args.tour:
        tour = next((t for t in SAMPLE_TOURS if args.tour.lower() in t['client'].lower()), None)
        if not tour:
            print(f"Tour not found. Available: {', '.join(t['client'] for t in SAMPLE_TOURS)}")
            sys.exit(1)
    elif args.random:
        tour = random.choice(SAMPLE_TOURS)
    else:
        tour = SAMPLE_TOURS[0]  # Default: Hugo Villar

    output_path = Path(args.output)
    generate_docx(tour, output_path)

    print(f"✓ Generated: {output_path}")
    print(f"  Client:   {tour['client']}, {tour['country']}")
    print(f"  Pax:      {tour['pax']}")
    print(f"  Days:     {tour['days']}")
    print(f"  Cities:   {', '.join(tour['cities'])}")
    print(f"")
    print(f"Test against your webhook:")
    print(f"  bash test-webhook.sh {output_path}")


if __name__ == '__main__':
    main()
